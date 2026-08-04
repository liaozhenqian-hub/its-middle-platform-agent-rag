from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("中台知识问答Agent系统学习与面试指南.docx")
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "59636E"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE = "F4F6F9"
WHITE = "FFFFFF"
RED = "9B1C1C"
GOLD = "7A5A00"
FONT_CN = "Microsoft YaHei"
FONT_LATIN = "Calibri"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size=None, bold=None, color=None, italic=None, mono=False) -> None:
    name = "Consolas" if mono else FONT_CN
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas" if mono else FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Consolas" if mono else FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph_runs(paragraph, size=10.5, color=NAVY, bold=False) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("中台知识问答 Agent 系统 | 学习与面试指南")
    set_run_font(hr, size=9, color=MUTED, bold=True)
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_title_page(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("中台知识问答 Agent 系统")
    set_run_font(r, size=28, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run("系统学习与面试指南")
    set_run_font(r, size=20, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(50)
    r = p.add_run("业务理解 · 架构设计 · RAG 检索 · 多智能体 · 工程化 · 面试表达")
    set_run_font(r, size=11.5, color=MUTED)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    values = [
        ("适用对象", "希望系统掌握本项目，并用于项目复盘或求职面试"),
        ("项目阶段", "具备 dev / 小范围试用能力，仍在做准确率与性能收敛"),
        ("技术基线", "FastAPI + OpenAI Agents SDK + LangGraph + PostgreSQL + pgvector"),
        ("文档日期", date.today().isoformat()),
    ]
    for row, (label, value) in zip(table.rows, values):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        style_paragraph_runs(row.cells[0].paragraphs[0], bold=True, color=DARK_BLUE)
        style_paragraph_runs(row.cells[1].paragraphs[0])
    set_table_geometry(table, [2700, 6660])
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("说明：本文按当前代码基线整理；规划中但尚未通过完整验证的能力会明确标注。")
    set_run_font(r, size=9.5, color=GOLD, italic=True)
    doc.add_page_break()


def add_heading(doc, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc, text: str, *, bold_prefix: str | None = None, color=NAVY) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, color=color)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, color=color)


def add_bullets(doc, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, color=NAVY)


def add_numbers(doc, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r, color=NAVY)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, LIGHT_BLUE)
        style_paragraph_runs(cell.paragraphs[0], size=9.5, bold=True, color=DARK_BLUE)
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].text = value
            style_paragraph_runs(row.cells[i].paragraphs[0], size=9.2)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_callout(doc, title: str, text: str, fill=PALE, color=DARK_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + "\n")
    set_run_font(r, size=10.5, bold=True, color=color)
    r = p.add_run(text)
    set_run_font(r, size=9.8, color=NAVY)
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_flow(doc, steps: list[str]) -> None:
    table = doc.add_table(rows=1, cols=len(steps))
    table.style = "Table Grid"
    width = 9360 // len(steps)
    widths = [width] * len(steps)
    widths[-1] += 9360 - sum(widths)
    for index, (cell, value) in enumerate(zip(table.rows[0].cells, steps)):
        cell.text = value
        set_cell_shading(cell, LIGHT_BLUE if index % 2 == 0 else PALE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph_runs(p, size=8.6, bold=True, color=DARK_BLUE)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_toc(doc: Document) -> None:
    add_heading(doc, "阅读导航", 1)
    add_para(doc, "建议先读第 1、3、4、7 章建立全局认识，再读第 17 章准备面试；需要深入某个专题时回到对应章节。")
    rows = [
        ["第一部分：为什么做", "1-2", "业务问题、能力边界与成熟度"],
        ["第二部分：系统怎么工作", "3-8", "总体架构、问答链路、多智能体、RAG、Bug、MCP"],
        ["第三部分：如何工程化", "9-14", "记忆、存储、质量、前端、安全、性能"],
        ["第四部分：如何继续建设", "15-16", "限制、上线建议、学习路线"],
        ["第五部分：如何用于面试", "17-19", "面试题、项目介绍模板、术语表"],
    ]
    add_table(doc, ["部分", "章节", "内容"], rows, [3100, 1000, 5260])
    add_callout(doc, "先记住一句话", "这是一个面向企业中台的证据驱动问答系统：用多智能体做领域路由，用混合 RAG 找到产品文档、代码和接口证据，用 LangGraph 约束 Bug 排查流程，再用引用门禁、记忆隔离和回归评测控制回答质量。")
    doc.add_page_break()


def build_content(doc: Document) -> None:
    add_heading(doc, "1. 系统解决的业务问题", 1)
    add_para(doc, "企业中台往往同时维护审批流、工作流、指标平台等多个复杂系统。其他部门的开发、产品和实施人员在对接时，经常遇到“文档在哪、接口怎么传、某字段代表什么、需求能否实现、线上报错如何定位”等问题。答案散落在产品文档、Git 分支、Swagger、指标 MCP 和日志平台中，依赖少数熟悉系统的人进行人工解释。")
    add_heading(doc, "1.1 典型痛点", 2)
    add_bullets(doc, [
        "知识分散：产品规则、代码实现、接口契约、运行日志彼此割裂，检索成本高。",
        "环境差异：develop 与 master 的实现可能不同，代码存在也不能等价为已经发布。",
        "问题表述不标准：用户会使用口语、简称或缺失关键信息，需要先理解意图再检索。",
        "回答容易过度推断：模型可能用通用经验补齐内部事实，产生“看起来合理但没有证据”的答案。",
        "Bug 排查链路长：需要环境、trace ID、日志信号、对应分支代码和接口文档共同支持。",
        "质量难持续验证：一次回答不错不代表长期稳定，需要真实问题集、事实约束和自动回归。",
    ])
    add_heading(doc, "1.2 系统提供的业务价值", 2)
    add_table(doc, ["用户场景", "系统行为", "预期价值"], [
        ["接口对接", "检索产品文档、代码与 Swagger，整理 URL、入参、出参和限制", "缩短跨部门沟通和联调时间"],
        ["产品使用", "按审批流、工作流或指标平台领域回答操作与规则", "降低重复咨询"],
        ["需求分析", "结合现有规则与代码，区分可行性、影响和未知项", "更早发现约束与改造成本"],
        ["Bug 定位", "按 trace 查询日志，提取异常信号，再关联对应分支代码", "把人工排障步骤固化为流程"],
        ["指标查询", "通过受控 MCP 查指标定义、应用候选和只读数据", "避免模型猜数据或绕过平台"],
    ], [1800, 4560, 3000])
    add_callout(doc, "核心定位", "系统不是通用聊天机器人，也不是自动修改生产系统的执行 Agent。它是一个面向企业中台知识、对接和故障分析的只读智能助手。", fill="FFF8E8", color=GOLD)

    add_heading(doc, "2. 能力边界与当前成熟度", 1)
    add_table(doc, ["能力", "当前状态", "边界"], [
        ["审批流 / 工作流问答", "已具备", "正确性受知识覆盖、分支差异和检索质量影响"],
        ["指标平台知识与实时查询", "已具备", "实时数据只能通过受控 MCP；写操作禁止"],
        ["代码 / 文档 / Swagger 检索", "已具备", "代码存在不代表部署；Swagger 缺失时不能确认完整契约"],
        ["LangGraph Bug 诊断", "已具备", "需要环境与原始 trace ID；无日志时不继续猜代码根因"],
        ["会话、长期与程序性记忆", "已具备基础能力", "必须按 owner、scope、space、domain 隔离"],
        ["质量采集与 Critical 回归", "框架已具备", "完整 5/10/30 回归仍需持续收敛"],
        ["生产多副本高可用", "尚未完成", "当前单 Worker，内存 BM25 和内置后台任务限制水平扩容"],
    ], [1900, 1800, 5660])
    add_callout(doc, "成熟度判断", "从代码和功能角度，系统已经具备 dev / 小范围试用的完整骨架；但不应宣称已经达到大规模商用成熟。当前 Critical 5 的已知基线为 3/5：或签、会签拒绝、管理员转办通过；实例详情的分支字段差异仍需稳定验证，加签缺少完整前后流转规则证据。", fill="FDECEC", color=RED)

    add_heading(doc, "3. 总体架构", 1)
    add_flow(doc, ["Web / API", "FastAPI", "意图路由", "领域专家", "证据工具", "模型回答"])
    add_para(doc, "系统采用分层设计，而不是让一个大模型自由决定所有动作。请求入口、身份与会话、意图路由、专家执行、检索、证据治理、回答生成和质量记录各自承担明确职责。")
    add_table(doc, ["层次", "关键职责", "主要技术"], [
        ["交互层", "网页聊天、SSE 流式响应、历史会话、引用详情、管理后台", "Vue 3、TypeScript、Vite"],
        ["API 层", "请求校验、生命周期、readiness、身份与质量采集", "FastAPI、Pydantic、Uvicorn"],
        ["Agent 层", "Manager、领域专家、工具调用、回答模板和证据门禁", "OpenAI Agents SDK"],
        ["流程层", "Bug 的暂停、恢复、固定节点和证据分级", "LangGraph、PostgreSQL Checkpointer"],
        ["检索层", "Query Rewrite、BM25、向量召回、融合和精排", "jieba、BM25Plus、pgvector、RRF、Qwen Rerank"],
        ["数据层", "会话、知识目录、质量、记忆、事件和向量", "PostgreSQL 15+、pgvector 0.8.x、Alembic"],
        ["外部工具层", "Git、Swagger、Grafana Loki、指标平台", "只读适配器、MCP、受控客户端"],
    ], [1500, 4380, 3480])
    add_heading(doc, "3.1 为什么不是一个万能 Agent", 2)
    add_bullets(doc, [
        "领域隔离：审批流、工作流和指标平台使用固定 domain 过滤，降低跨领域误召回。",
        "权限隔离：指标实时数据、Grafana 和 Swagger 由服务端控制参数，模型不能任意构造访问。",
        "成本控制：高置信度单领域问题直接进入专家，减少 Manager 的重复模型调用。",
        "可测试性：确定性路由、检索预算和证据规则可以写成单元测试和硬门禁。",
        "可降级性：Bug Graph、MCP 或 Rerank 不可用时，其他知识问答能力仍可按规则降级。",
    ])

    add_heading(doc, "4. 一次问答的完整链路", 1)
    add_numbers(doc, [
        "入口接收问题：网页调用 JSON 或 SSE 接口；服务端解析会话、用户身份、渠道和当前上下文。",
        "请求守卫：识别越界问题、敏感信息、明确写操作和取消指令，禁止无关闲聊或泄露凭证。",
        "意图路由：先用确定性规则识别领域与任务类型；规则无法判断时才用轻量模型补充分类。",
        "执行策略：明确单领域问题直接调用对应专家；跨领域问题由 Manager 调多个专家后汇总；Bug 进入 LangGraph。",
        "受控证据收集：根据任务类型选择产品文档、代码、Swagger 或 MCP，并限制重复查询与调用预算。",
        "混合检索：Query Rewrite 后并行运行 BM25 和 pgvector，RRF 融合候选，再由 Rerank 精排。",
        "证据门禁：过滤弱相关引用；没有证据时不得确认内部事实，但可以明确标注未知并给出待验证建议。",
        "答案生成：按 api_contract、how_to、code_lookup、requirement_analysis 等固定结构输出。",
        "流式返回：SSE 先返回阶段事件，再输出文本增量与最终结构化结果；页面展示引用和反馈入口。",
        "异步质量记录：保存问题、回答、路由、工具、引用、耗时、反馈和问题标签，用于后续回归。",
    ])
    add_callout(doc, "关键原则", "大模型负责理解口语和组织语言，代码负责权限、状态、参数、分支、预算和证据边界。可靠性来自二者分工，而不是单纯换一个更强模型。")

    add_heading(doc, "5. 多智能体与路由设计", 1)
    add_heading(doc, "5.1 Agent 拓扑", 2)
    add_table(doc, ["角色", "负责什么", "不负责什么"], [
        ["Manager", "跨领域协调、最终汇总、范围提示", "不凭模型记忆回答内部事实"],
        ["审批流专家", "审批规则、接口、代码和产品文档", "不检索工作流或指标平台数据"],
        ["工作流专家", "节点、变量、连接器、接口与代码", "不擅自查询日志或执行连接器"],
        ["指标平台专家", "指标知识、SDK、MCP 只读查询", "不直接访问数据库或猜指标应用"],
        ["Bug 诊断专家", "日志、代码、契约关联和诊断报告", "缺环境/trace 时不猜根因"],
    ], [1700, 4200, 3460])
    add_heading(doc, "5.2 ReAct 用在哪里", 2)
    add_para(doc, "系统使用了 ReAct 的“理解—行动—观察—回答”思想，但并非全链路自由 ReAct。领域专家可以在工具约束内调用证据工具；Bug 助手则使用 LangGraph 固定节点，将关键步骤变成可恢复状态机。")
    add_table(doc, ["模式", "优点", "风险", "本项目策略"], [
        ["自由 ReAct", "灵活处理开放问题", "循环调用、成本不可控、难测试", "只在受控专家工具范围内有限使用"],
        ["确定性工作流", "稳定、可审计、可暂停恢复", "适应性略低", "Bug 诊断、MCP 查询状态采用此模式"],
        ["规则 + LLM 路由", "常见问题快，口语问题也能覆盖", "规则需维护，模型可能误判", "规则优先，Flash 仅处理未决问题"],
    ], [1600, 2500, 2580, 2680])
    add_heading(doc, "5.3 任务类型决定回答模板", 2)
    add_table(doc, ["task_type", "回答重点"], [
        ["api_contract", "结论；接口地址与方法；请求/响应字段；必填与校验；证据；未确认事项"],
        ["how_to", "结论；适用场景；操作步骤；限制；证据；未确认事项"],
        ["code_lookup", "结论；代码位置；实现说明；验证方式；证据；未确认事项"],
        ["requirement_analysis", "结论；现有能力；可行性与影响；实施建议；证据；未确认事项"],
        ["metric_query", "指标候选确认；只读查询结果；口径与时间范围；证据"],
        ["bug", "问题摘要；日志事实；代码关联；原因和置信度；修复与验证"],
    ], [1900, 7460])

    add_heading(doc, "6. RAG 检索链路", 1)
    add_flow(doc, ["原问题", "Query Rewrite", "BM25 + Vector", "RRF", "Rerank", "强证据"])
    add_heading(doc, "6.1 Query Rewrite", 2)
    add_para(doc, "用户问题往往包含口语、路径、字段名和上下文。改写模型输出 retrieval_query、关键词、领域候选和任务类型。向量检索使用更自然的语义查询，BM25 使用保序去重后的精确关键词。若改写失败，系统回退到原问题，不让增强步骤阻断检索。")
    add_heading(doc, "6.2 Metadata 过滤", 2)
    add_para(doc, "在召回前先限制 app_id、domain、source_type、branch 等可信字段。例如审批流 develop 代码查询只允许进入 approval-flow + code + develop 的候选空间。它既提升速度，也避免跨领域、跨分支证据污染。")
    add_heading(doc, "6.3 BM25 与 pgvector", 2)
    add_table(doc, ["召回方式", "擅长", "不擅长", "项目中的角色"], [
        ["BM25", "接口路径、类名、字段名、错误码等精确词", "同义表达和语义相似", "使用 heading 与 bm25_keywords 的内存索引做关键词召回"],
        ["pgvector", "口语问题、同义表达、语义相关内容", "精确标识符可能被稀释", "使用 1024 维向量和 cosine distance 做语义召回"],
    ], [1500, 2700, 2300, 2860])
    add_para(doc, "BM25 冷启动时只读取轻量 metadata，而不把所有长正文加载进内存；选出 Top K 后再按 ID 回填正文。当前知识向量规模约 4 万条，产品文档曾完成 268 条领域映射回填。")
    add_heading(doc, "6.4 RRF 融合", 2)
    add_para(doc, "BM25 分数越大越好，向量距离越小越好，二者量纲不同，不能直接相加。RRF（Reciprocal Rank Fusion）只利用名次：某条证据在多个召回列表中排名越靠前，融合分越高。常见形式为 1 / (k + rank)，本项目默认思想是先合并去重，再按融合分得到稳定候选。")
    add_heading(doc, "6.5 Rerank 精排", 2)
    add_para(doc, "Reranker 同时阅读问题与候选的标题、关键词和正文，判断候选是否真正回答问题。它比向量相似度更适合做最终相关性排序，但需要额外网络调用。因此系统设置超时、熔断和降级：Rerank 不可用时退回 RRF，不让整个问答失败。")
    add_heading(doc, "6.6 引用与证据门禁", 2)
    add_bullets(doc, [
        "默认只展示 3 至 5 个最强证据，弱相关内容不公开。",
        "代码中的类、方法或 Controller 可以证明该分支存在实现；不能证明已经部署。",
        "缺 Swagger 只影响完整接口契约确认，不能否定已有代码事实。",
        "缺发布记录只影响部署状态确认，不能把正确代码答案整体标记为未找到。",
        "用户可见来源使用中文可读名称，不展示内部 chunk ID 或 source ID。",
        "日志和 MCP 只展示受控摘要，不返回原始日志或完整工具输出。",
    ])

    add_heading(doc, "7. LangGraph Bug 诊断", 1)
    add_para(doc, "Bug 是最需要流程约束的场景。系统要求用户提供环境与原始 trace ID，按固定节点查询日志、抽取事实、关联代码、补充契约、评定证据等级并生成报告。")
    add_flow(doc, ["理解输入", "校验 / 暂停", "查日志", "提取信号", "查代码", "证据分级", "报告"])
    add_heading(doc, "7.1 环境与分支映射", 2)
    add_table(doc, ["用户表达", "日志环境", "代码分支"], [
        ["开发、develop、dev", "develop", "develop"],
        ["测试、test", "test", "develop（当前无独立 test 代码分支）"],
        ["线上、生产、prod、production", "prod", "master"],
    ], [3000, 2200, 4160])
    add_heading(doc, "7.2 为什么使用 interrupt", 2)
    add_para(doc, "缺环境或 trace ID 时，Graph 使用 interrupt 持久化暂停，而不是返回一个没有状态的追问。用户在同一 conversation 中补充信息后，通过 Command(resume=...) 恢复；24 小时后过期。暂停状态只保存结构化字段和引用标识，不保存日志正文、代码正文、凭证、Prompt 或完整模型响应。")
    add_heading(doc, "7.3 日志优先与证据等级", 2)
    add_table(doc, ["等级", "含义", "允许输出"], [
        ["none", "没有可用日志或代码证据", "只说明无法确认及缺失信息"],
        ["log_only", "日志确认了异常，但未关联可靠代码", "日志事实与排查建议，不声明代码根因"],
        ["correlated", "日志信号与对应分支代码相关联", "可能的代码根因、置信度、修复和验证"],
        ["contract_supported", "在 correlated 基础上还有 Swagger / 产品文档支持", "更完整的接口和规则分析"],
    ], [1500, 4300, 3560])
    add_para(doc, "Grafana 查询参数由服务端固定控制，默认按 trace ID 查询最近 24 小时；模型不能修改 URL、LogQL、datasource、namespace 或代码分支。网络、超时和 5xx 可有限重试，认证和参数错误不重试。零日志时立即停止代码检索，避免无日志依据的“猜根因”。")
    add_callout(doc, "面试中的亮点", "LangGraph 的价值不是“画了一张图”，而是把必须按顺序执行、需要暂停恢复、必须审计且不能让模型自由修改参数的诊断流程，变成显式状态机。")

    add_heading(doc, "8. 指标 MCP 状态机与缓存", 1)
    add_para(doc, "指标平台既有知识问答，也有实时数据查询。产品说明、SDK、配置和代码位置走 RAG；实际指标值与 SQL 来源只能走 MCP。这样避免大模型把历史文档中的示例数据当成实时结果。")
    add_numbers(doc, [
        "调用 metricMcpInfo 获取平台规则。",
        "使用 searchBizMetric 查找业务指标候选。",
        "使用 searchMetricApp 确认指标应用。",
        "若用户未明确选择应用，状态机返回 clarification_required 并停止，不自行猜选。",
        "确认后由 prepare_metric_query 生成受控查询，再调用 allowlist 中的只读查询工具。",
        "同一标准化查询复用代码级缓存；状态和候选写入会话上下文，避免重复调用。",
    ])
    add_bullets(doc, [
        "服务端只暴露 allowlist 工具，MCP 新增工具不会自动获得 Agent 权限。",
        "同一轮每个发现工具最多调用一次，已有候选必须复用。",
        "MCP 不可用时知识库问答仍可用，实时数据明确提示不可用。",
        "指标 MCP 的完整输出不会直接返回用户，也不会被当作模型自由执行入口。",
    ])

    add_heading(doc, "9. 记忆系统", 1)
    add_para(doc, "知识库回答“组织已经知道什么”，会话上下文回答“这一轮在聊什么”，长期记忆回答“这个用户或流程过去确定了什么”。三者目的不同，不能混为一个向量库。")
    add_table(doc, ["记忆类型", "示例", "生命周期与用途"], [
        ["会话记忆", "当前问题、追问、已选环境", "随 conversation 使用，支持多轮上下文"],
        ["摘要记忆", "长对话的压缩摘要", "减少上下文 token，保留主线"],
        ["个人偏好 / 上下文", "常用环境、回答偏好、所属系统", "候选经确认；普通类型可在 24 小时后自动确认"],
        ["实体记忆", "系统、应用、接口、角色之间的关系", "用于补充稳定实体上下文，需避免错误关系扩散"],
        ["Bug incident", "某类故障的症状、根因和验证结果", "帮助关联重复事件，但不能替代本次日志证据"],
        ["程序性记忆", "某类排障步骤或成功工具序列", "为 Bug Graph 提供流程建议，初期可观察而不自动执行"],
        ["领域记忆", "经治理后可复用的组织经验", "必须从个人记忆晋升并经过管理员审核"],
    ], [1800, 3100, 4460])
    add_heading(doc, "9.1 隔离模型", 2)
    add_para(doc, "个人记忆查询必须同时包含 collection、owner_id、scope_type 和 space_id，并按 domain 限制。匿名网页用户可通过稳定 Cookie 身份区分；飞书用户使用平台 open ID；Codex 等机器调用使用独立渠道身份。个人记忆由本人确认、驳回和删除，管理员只治理领域晋升，不应随意干预个人偏好。")
    add_callout(doc, "记忆的安全边界", "记忆只能帮助理解偏好和复用已确认事实，不能提升某条事实的证据等级。Bug 诊断即使召回历史 incident，也必须查询本次 trace 日志。", fill="FFF8E8", color=GOLD)

    add_heading(doc, "10. PostgreSQL 与 pgvector 持久化", 1)
    add_para(doc, "系统已从本地 SQLite + Chroma 逐步迁移为 PostgreSQL + pgvector。关系数据与向量数据使用同一个 PostgreSQL 基础设施，但通过不同表与仓储接口管理。")
    add_table(doc, ["内容", "当前存储", "设计要点"], [
        ["知识目录、任务、会话、质量、记忆、认证、飞书事件", "PostgreSQL", "SQLAlchemy Core、事务、JSONB、TIMESTAMPTZ"],
        ["知识向量与个人记忆向量", "pgvector", "统一 vector_entries，1024 维 cosine，collection + owner/scope 隔离"],
        ["Bug Graph checkpoint", "PostgreSQL Saver", "由 LangGraph 官方 Checkpointer 管理"],
        ["Schema 演进", "Alembic", "基线、升级和回滚；Graph 内部表不重复定义"],
        ["关键词索引", "进程内存", "BM25 冷启动预热，当前尚未迁移为 PostgreSQL 全文检索"],
    ], [2500, 2100, 4760])
    add_heading(doc, "10.1 为什么迁移", 2)
    add_bullets(doc, [
        "PostgreSQL 更适合 dev / 生产的并发、备份、审计和统一运维。",
        "pgvector 让 metadata 过滤与向量检索处于同一数据库事务和查询体系。",
        "关系仓储与向量仓储都采用 provider 工厂，API 与业务层不直接判断数据库类型。",
        "任务领取可使用 FOR UPDATE SKIP LOCKED，减少多任务重复领取。",
    ])
    add_heading(doc, "10.2 为什么迁移后可能变慢", 2)
    add_bullets(doc, [
        "本地 Chroma / SQLite 没有网络往返，dev PostgreSQL 经过 Telepresence 时网络延迟更明显。",
        "BM25 首次构建需要从 pgvector 读取 metadata；若未预热，首个问题会承担冷启动成本。",
        "过滤字段映射错误会扩大候选集，既降低准确率也增加查询时间。",
        "Rerank 和模型仍是外部网络调用，数据库并非全部延迟来源。",
        "需要通过 spans 分解 database、rewrite、BM25、vector、rerank、LLM，而不是凭总耗时猜测。",
    ])

    add_heading(doc, "11. 质量数据与 Critical 回归", 1)
    add_para(doc, "质量系统的目标是把真实用户问题变成可重复测试，而不是只看点赞。每轮记录渠道、领域、任务类型、模型、路由、工具、引用、耗时和反馈；质量 span 记录阶段耗时与 token，但不保存模型思维链、Prompt 或敏感正文。")
    add_heading(doc, "11.1 EvalCase 的约束", 2)
    add_table(doc, ["字段", "作用"], [
        ["turns", "支持多轮补充、取消、切换话题等场景"],
        ["required_facts", "回答必须覆盖的业务事实"],
        ["forbidden_facts", "禁止泄露或禁止声称的内容"],
        ["required_tools / citations", "确保走对专家并获得必要证据"],
        ["max_latency / max_tool_calls", "约束用户体验和 Agent 循环"],
        ["case version / snapshot", "让历史结果能还原当时模型、知识版本和配置"],
    ], [2800, 6560])
    add_heading(doc, "11.2 两层评分", 2)
    add_numbers(doc, [
        "硬门禁先检查状态、路由、工具、引用、安全、敏感信息、工具数、延迟和部署措辞。硬门禁失败时不浪费语义裁判调用。",
        "通过硬门禁后，由 DeepSeek 语义裁判评估相关性、事实正确性、证据支持、矛盾、未知项校准和可执行性。",
        "裁判输出必须通过 JSON 校验；修复一次仍失败则标记 judge_error，不能默认通过。",
    ])
    add_heading(doc, "11.3 Critical 集是什么", 2)
    add_para(doc, "Critical 不是随机问句，而是 30 条高风险业务回归，覆盖审批流接口与规则、工作流节点与变量、指标 MCP、安全拒绝、分支/部署措辞和无 Swagger 校准。每条都应有必需事实、禁止事实、引用类型和工具预算。")
    add_callout(doc, "当前真实状态", "不能用“测试框架已存在”代替“全部用例已通过”。当前已知 Critical 5 为 3/5，完整 5/10/30 尚未形成全部通过的发布门禁结果。因此扩大用户范围前，应先修复证据缺口与分支差异，再逐批回归。", fill="FDECEC", color=RED)

    add_heading(doc, "12. 前端、流式输出与引用详情", 1)
    add_table(doc, ["能力", "实现思路", "常见故障"], [
        ["SSE 流式聊天", "服务端发送阶段事件、文本 delta 和最终结果", "底层 Agent 只返回最终文本时，前端仍会感觉一次性输出"],
        ["历史会话", "按 owner 与 conversation scope 查询，刷新恢复当前 conversation", "身份 Cookie 变化或 owner 映射不一致会看不到历史"],
        ["引用面板", "先展示命中章节，点击后按需加载全文或受控详情", "source metadata 缺 URL、文档被停用或字段映射错误"],
        ["反馈", "点赞/点踩和原因进入质量库", "反馈覆盖低时需结合重问、纠正、超时等系统 annotation"],
    ], [1800, 4160, 3400])
    add_para(doc, "引用详情应面向用户显示可读标题、文档 URL、代码路径、分支、symbol、行号和 Git permalink。内部 chunk ID 只用于排查，不能成为页面上的“证据名称”。产品文档默认展示命中章节，用户点击“查看全文”后再分页读取整个文件，避免首屏过重。")

    add_heading(doc, "13. 安全与权限边界", 1)
    add_bullets(doc, [
        "客户端知识问答只读；不得通过 Agent 修改代码、数据库、知识源或调用 Swagger 中的业务写接口。",
        "凭证只从服务端环境读取，不写入日志、Graph checkpoint、Prompt、回答或迁移中间文件。",
        "Grafana、GitLab、MCP、Swagger 都使用固定客户端、allowlist 和参数校验，模型只提供受控业务输入。",
        "个人记忆按用户和空间隔离；领域记忆需要治理晋升，避免把个人错误经验传播给所有人。",
        "管理员接口使用独立认证；网页可选飞书登录或匿名 Cookie，生产应在网关进一步做身份与访问控制。",
        "日志和 MCP 只展示脱敏摘要；原始日志、Authorization、API Key、Embedding 和工具完整输出均不公开。",
    ])
    add_heading(doc, "13.1 Prompt Guardrail 不是唯一防线", 2)
    add_para(doc, "提示词可以告诉模型“不能做什么”，但真正的安全边界必须在代码中：工具 allowlist、只读客户端、固定查询范围、字段白名单、仓储隔离、输出脱敏和审计。只依赖 Prompt 会受到提示注入、模型漂移和供应商差异影响。")

    add_heading(doc, "14. 性能瓶颈与优化方法", 1)
    add_table(doc, ["阶段", "可能耗时", "优化方法"], [
        ["意图与改写", "一次或多次 LLM 往返", "规则优先；高置信度单领域直达专家；并发无依赖调用"],
        ["BM25", "冷启动读 metadata 和建索引", "服务启动预热；按 domain 缓存；stale-while-refresh"],
        ["pgvector", "网络、过滤、HNSW 参数", "独立过滤列与组合索引；连接池；ANALYZE；定位 Telepresence 延迟"],
        ["Rerank", "外部模型调用", "缩小候选；超时与熔断；结果缓存；失败回退 RRF"],
        ["Agent 工具", "重复同义查询或循环", "标准化查询去重；单领域底层检索最多 4 次；聚合证据工具"],
        ["最终模型", "长上下文、思考模型", "Flash 用于路由/常规回答；Pro 只用于高价值无工具综合"],
        ["前端", "长时间无可见反馈", "真实 token 流或阶段事件；超时提示与取消"],
    ], [1500, 2760, 5100])
    add_heading(doc, "14.1 正确的排查方法", 2)
    add_para(doc, "先用 quality_spans 得到阶段瀑布，再针对最大项优化。若 database 只占 200 ms，而最终 LLM 占 20 s，迁移数据库不会解决问题；若首问 BM25 10 s、后续 100 ms，则应做启动预热而不是更换模型。性能判断必须基于分段数据。")

    add_heading(doc, "15. 当前限制与上线建议", 1)
    add_heading(doc, "15.1 当前限制", 2)
    add_bullets(doc, [
        "单 Uvicorn Worker：BM25 内存索引、内置 Worker 和 Git mirror 尚未完全拆分。",
        "Critical 尚未 30/30：存在分支字段差异和产品规则证据缺口。",
        "外部模型、Embedding、Rerank、Grafana、MCP 都会引入网络依赖和降级场景。",
        "BM25 仍为进程内状态，多副本需要统一刷新机制或改用外部全文检索。",
        "长期记忆仍需持续评估错误记忆、冲突处理、删除一致性与隐私治理。",
    ])
    add_heading(doc, "15.2 dev 上线门槛", 2)
    add_numbers(doc, [
        "固定内网域名与回调地址，使用 Kubernetes Secret 注入配置，禁止把 .env 打进镜像。",
        "使用持久卷保存必要文件；PostgreSQL、pgvector、模型、Grafana 和 MCP 均加入 readiness。",
        "当前保持单副本，确保 Source Worker、评测 Worker、记忆 Worker 和飞书长连接不重复启动。",
        "启动时完成 Alembic upgrade、LangGraph saver setup 和 BM25 预热；失败组件按关键性决定阻断或降级。",
        "先跑 smoke 和 Critical 5，再依次 10、30；结果未全部通过前只开放小范围用户。",
        "监控 P50/P90、零引用率、工具数、错误路由、重复追问、超时和外部依赖可用性。",
        "准备 PostgreSQL 备份、应用回滚、旧向量只读保留和知识同步恢复方案。",
    ])

    add_heading(doc, "16. 推荐学习路线", 1)
    add_table(doc, ["阶段", "学习目标", "建议实践"], [
        ["第 1 周：业务与入口", "讲清审批流、工作流、指标平台痛点与一次请求链路", "画架构图；手工跟踪一个 api_contract 请求"],
        ["第 2 周：RAG", "掌握切分、Embedding、BM25、向量、RRF、Rerank、证据门禁", "对同一问题比较纯向量、纯 BM25 与混合结果"],
        ["第 3 周：Agent", "理解路由、专家、工具、ReAct 与确定性工作流的边界", "增加一个只读专家工具并写预算测试"],
        ["第 4 周：LangGraph", "理解 state、node、edge、interrupt、checkpoint、resume", "用 fake Grafana 跑一轮暂停与恢复"],
        ["第 5 周：存储与记忆", "掌握 PostgreSQL、pgvector、作用域隔离和迁移", "解释 vector_entries 的过滤列和 HNSW"],
        ["第 6 周：评测与性能", "能用事实约束、引用和 spans 判断质量", "修复一个 Critical 失败并做前后对比"],
        ["第 7 周：上线", "理解单 Worker 限制、降级、Secret、备份和可观测性", "完成一份 dev 发布检查单"],
    ], [1700, 3900, 3760])

    add_heading(doc, "16.1 源码阅读地图", 2)
    add_para(doc, "建议按“入口 -> 编排 -> 检索 -> 专项流程 -> 持久化 -> 评测”的方向阅读，不要从工具类随机跳转。以下路径均相对于项目根目录。")
    add_table(doc, ["阅读顺序", "核心文件 / 目录", "重点问题"], [
        ["1. 启动与生命周期", "knowledge/api/app.py", "组件如何初始化、降级、关闭；readiness 如何聚合"],
        ["2. 配置", "knowledge/config/settings.py", "provider、模型、预算、超时和安全配置如何校验"],
        ["3. Agent 拓扑", "knowledge/agent_runtime/agent_factory.py", "Manager、专家、工具和提示词如何组装"],
        ["4. 意图路由", "intent_router.py / hybrid_intent_router.py", "规则何时直接命中，何时调用 Flash fallback"],
        ["5. Agent 服务", "knowledge/agent_runtime/service.py", "请求如何执行、证据如何门禁、结果如何归一化"],
        ["6. 运行上下文", "knowledge/agent_runtime/context.py", "工具预算、重复查询、引用选择和 response_mode"],
        ["7. 聚合证据工具", "knowledge/agent_runtime/rag_tools.py", "文档、代码和 Swagger 如何并发且受预算控制"],
        ["8. 检索流水线", "knowledge/services/multi_route_retrieval_service.py", "rewrite、BM25、vector、merge/rerank 的时序"],
        ["9. 关键词召回", "knowledge/services/keyword_retrieval_service.py", "轻量 metadata、字段权重、过滤和正文回填"],
        ["10. 融合精排", "hybrid_rerank_service.py / qwen_rerank_service.py", "RRF 去重、熔断、降级和候选组织"],
        ["11. Bug Graph", "knowledge/bug_graph/service.py", "节点、条件边、interrupt、resume 和证据分级"],
        ["12. 向量仓储", "knowledge/repositories/vector_store_repository.py", "provider 协议、过滤、分页和 pgvector 语义"],
        ["13. 记忆", "knowledge/memory/", "候选提取、确认、摘要、实体和程序性记忆"],
        ["14. 质量评测", "knowledge/quality/", "数据采集、硬门禁、语义裁判和 Critical 定义"],
        ["15. 前端", "web/src/", "SSE、会话恢复、引用详情、记忆和质量管理页面"],
    ], [1100, 3550, 4710])
    add_heading(doc, "16.2 建议断点与观测字段", 2)
    add_bullets(doc, [
        "在 API chat/stream 入口观察 conversation_id、owner、channel；不要记录问题正文到调试截图。",
        "在 HybridDomainIntentRouter.route 观察 domains、intent、task_type、route_source 和 duration_ms。",
        "在 AgentRunContext 的工具登记处观察 normalized query、实际调用数、重复拦截和 citations。",
        "在 MultiRouteRetrievalService 观察 rewrite、keyword/vector 候选 ID、各阶段 timing 和 rerank_applied。",
        "在 Bug Graph 每个 node 观察状态字段变化，不查看或落盘原始日志和代码正文。",
        "在 quality_spans 比较 route、rewrite、vector、rerank、tool、LLM 的耗时，定位真实瓶颈。",
    ])

    add_heading(doc, "16.3 动手实验清单", 2)
    add_table(doc, ["实验", "操作", "验收结果"], [
        ["1. 路由", "分别提问审批流、工作流、指标平台和无关问题", "领域、任务类型和拒答边界正确"],
        ["2. 分支", "用同一代码问题分别指定 develop 与 prod", "引用分支不同，且不把代码存在说成已部署"],
        ["3. 混合召回", "用接口路径和口语描述查询同一能力", "BM25 与向量各有优势，RRF 后目标证据进入前列"],
        ["4. Rerank 降级", "在测试中模拟 Rerank 超时", "请求成功并退回 RRF，证据措辞不过度确认"],
        ["5. 引用门禁", "构造零命中与弱相关问题", "内部事实不被确认，弱引用不展示"],
        ["6. 工具预算", "让模型用同义词重复搜索", "标准化查询只实际执行一次，总调用不超过预算"],
        ["7. Bug interrupt", "先只给报错，再补环境和 trace", "Graph 暂停、同会话恢复，且只问缺失关键字段"],
        ["8. Bug 无日志", "fake Grafana 返回零条", "流程停止，不继续检索代码或编造根因"],
        ["9. MCP 澄清", "查询一个对应多个指标应用的指标", "返回候选并等待确认，不自动选应用"],
        ["10. 记忆隔离", "两个 owner 写入不同偏好后互相查询", "只召回各自 owner/scope/space 的内容"],
        ["11. SSE", "记录请求、首事件、首文本和完成时间", "区分阶段流式与真实文本 delta"],
        ["12. Critical", "先跑 5 条并阅读每条 hard gate", "能解释失败属于证据、路由、事实、延迟还是工具预算"],
    ], [1200, 4440, 3720])
    add_callout(doc, "学习产出建议", "每完成一个实验，保留一页笔记：问题、预期、实际路由、工具、证据、耗时、失败原因和改进。六周后这些材料可以直接转化为面试中的项目复盘。")

    add_heading(doc, "17. 面试问题与参考回答", 1)
    add_para(doc, "下面的回答是“要点”，面试时应先给结论，再结合项目中的真实取舍说明，不要背成定义。")
    questions = [
        ("Q1：请介绍一下这个项目。", "这是一个面向企业中台对接与故障分析的证据驱动 Agent。它把审批流、工作流、指标平台的产品文档、代码、Swagger、日志和指标 MCP 统一到一个问答入口。系统用规则与模型做领域路由，用 BM25 + pgvector + RRF + Rerank 做混合检索，用引用门禁限制无证据结论；复杂 Bug 使用 LangGraph 固化日志到代码的诊断流程。"),
        ("Q2：为什么纯向量检索不够？", "企业问题包含大量接口路径、字段、类名和错误码。向量适合语义相似，但可能弱化精确标识符；BM25 擅长字面匹配。因此并行召回，再用 RRF 融合不同量纲的排序，最后用 Rerank 判断候选是否真正回答问题。"),
        ("Q3：RRF 为什么比直接加分更合理？", "BM25 是相关性分数，向量通常是距离，量纲和方向都不同。直接加权需要脆弱的归一化。RRF 只依赖各列表排名，公式通常是 1/(k+rank)，对不同召回器更稳健，也便于降级。"),
        ("Q4：Rerank 与向量检索有什么区别？", "向量检索通常用双塔模型，query 和文档预先独立编码，速度快，适合大规模召回；Rerank 同时阅读 query 与候选文本，交互更充分，精度高但成本大，所以只处理小候选集。"),
        ("Q5：如何降低 RAG 幻觉？", "固定领域和分支过滤；强证据阈值；代码、Swagger、发布记录分别证明不同事实；零引用不能确认内部事实；答案明确区分已确认、推断和未知；Critical 用例检查必需事实、禁止事实和引用。"),
        ("Q6：为什么要做多智能体？", "不是为了增加 Agent 数量，而是为了隔离领域知识、工具权限和回答规则。高置信度单领域请求直接进专家，避免 Manager 往返；只有跨领域才汇总，从而控制成本和错误路由。"),
        ("Q7：这个项目算 ReAct 吗？", "部分算。领域专家有受控的思考—工具—观察链路，但不允许无限自由循环；Bug 与指标查询采用确定性状态机。项目的核心取舍是把开放理解交给 LLM，把权限、状态、预算与关键顺序交给代码。"),
        ("Q8：为什么 Bug 用 LangGraph？", "Bug 需要强顺序、缺字段暂停、跨轮恢复、状态审计和严格证据分级。LangGraph 的 state、interrupt、checkpoint 和 Command(resume) 比自由 Agent 更适合。它还保证模型不能改日志范围和代码分支。"),
        ("Q9：如何防止 Agent 重复调用工具？", "在 AgentRunContext 中按标准化 query + domain + source type 记账，限制相同查询次数和总预算；用聚合 evidence tool 统一并行检索代码、文档和 Swagger；MCP 使用会话状态和缓存复用候选。"),
        ("Q10：为什么从 Chroma 迁到 pgvector？", "主要是统一 dev/生产运维、metadata 过滤、事务与备份，并为后续并发提供 PostgreSQL 能力。迁移复用已有 embedding，不重新调用模型；先关系库切换，再 shadow 对比，最后切正式向量读取。"),
        ("Q11：pgvector 如何设计表和索引？", "统一 vector_entries，以 collection_name + id 为主键；正文、metadata 和 1024 维 embedding 共存；app、domain、source、branch、owner、scope、space 等高频过滤字段独立成列；向量使用 cosine HNSW，metadata 用 GIN 和组合索引。"),
        ("Q12：HNSW 的优缺点是什么？", "查询快、召回率高，适合在线近似最近邻；代价是索引占用大、构建和写入成本高。批量迁移时先导入再建 HNSW，并通过 ef_search 在召回率与延迟间权衡。"),
        ("Q13：BM25 为什么放内存？", "当前规模约 4 万，字段化轻量 metadata 可以快速构建并获得低延迟；缺点是每个进程各有一份且刷新一致性复杂。因此当前保持单 Worker，未来多副本可迁 PostgreSQL FTS / OpenSearch 或建立统一索引服务。"),
        ("Q14：如何评估问答准确率？", "不能只看是否有引用。硬门禁检查路由、工具、引用、安全、工具数和延迟；语义裁判检查事实、引用支持、矛盾和可执行性；高风险 Critical 用例还配置 required_facts 与 forbidden_facts，并保留知识和模型快照。"),
        ("Q15：LLM-as-a-Judge 有什么风险？", "裁判也会漂移、偏好长答案或受提示影响。应先跑确定性硬规则，裁判只看脱敏限长证据；固定模型和 rubric，校验 JSON，失败关闭；边界分数人工复核，并用少量人工金标校准裁判。"),
        ("Q16：长期记忆与知识库有什么区别？", "知识库是组织级可信资料；记忆是用户或会话相关信息。记忆必须按 owner/scope/space/domain 隔离，候选需确认，个人经验要经过治理才能晋升领域记忆；记忆不能替代本轮业务证据。"),
        ("Q17：如何保证不同用户记忆不串？", "身份解析后，每次写入和检索都强制带 owner_id、scope_type、space_id 和 collection；仓储与向量过滤都执行同样约束，并用匿名、飞书、Codex 等多渠道隔离测试验证。"),
        ("Q18：为什么网页看起来不是流式？", "SSE 端点存在不等于底层真正 token 流。如果 Agent SDK 或模型适配器只在完成后返回整段文本，前端只能一次展示。应区分阶段事件与文本 delta，并测量首字节和首 token 时间。"),
        ("Q19：系统慢时怎么定位？", "通过 quality_spans 分解 route、rewrite、database、BM25、vector、rerank、tool、final LLM。首问慢后续快通常是 BM25 冷启动；所有请求都慢可能是外部模型或 Telepresence；重复工具多则是 Agent 规划问题。"),
        ("Q20：代码存在为什么不能说已发布？", "Git 分支证据只能证明仓库中存在实现，部署需要 CI/CD、版本或环境探针证据。系统把事实证明范围拆开，避免缺发布记录时否定代码，也避免把代码误当部署状态。"),
        ("Q21：如何处理外部服务故障？", "组件按关键性降级：Rerank 失败退 RRF，MCP 失败保留知识问答，Bug Graph 或 Grafana 不可用时明确状态；readiness 展示组件状态。网络/5xx 可重试，认证/参数错误不重试，避免无意义放大流量。"),
        ("Q22：Prompt 注入怎么防？", "Prompt 只是一层。核心是服务端工具 allowlist、只读权限、固定参数、输入白名单、输出脱敏、Secret 隔离和审计。检索文档中的指令按数据处理，不能覆盖系统策略。"),
        ("Q23：如果做成多副本，要改什么？", "先把 Source、Eval、Memory Worker 拆成独立进程并加分布式任务领取；BM25 改共享索引服务或外部全文检索；飞书长连接单独部署；会话锁和缓存迁 Redis；文件与 Git mirror 放共享存储或对象存储；再做水平扩容。"),
        ("Q24：这个项目最大的技术取舍是什么？", "不是追求完全自主 Agent，而是在灵活性和可控性之间分层：口语理解与答案表达由模型完成，内部事实必须经过检索，Bug 和 MCP 用状态机，权限和证据由代码保证。这让系统更适合企业内部场景。"),
        ("Q25：目前最大的不足是什么？", "关键领域证据覆盖仍不完整，Critical 尚未全绿；单 Worker 与内存 BM25 限制扩容；外部模型链路影响 P90；记忆治理仍需长期验证。面试中应如实说明这些限制和下一步，而不是宣称完全商用成熟。"),
    ]
    for question, answer in questions:
        add_heading(doc, question, 3)
        add_para(doc, answer)

    add_heading(doc, "18. 项目介绍模板", 1)
    add_heading(doc, "18.1 30 秒版本", 2)
    add_callout(doc, "示例", "我做了一个企业中台知识问答 Agent，主要解决审批流、工作流和指标平台的接口对接、产品咨询、需求分析与 Bug 定位。系统用多智能体做领域隔离，用 BM25 和 pgvector 混合检索文档、代码与 Swagger，再通过 RRF、Rerank 和引用门禁控制回答。复杂 Bug 用 LangGraph 查询日志并关联对应分支代码，同时有长期记忆和 Critical 回归评测。")
    add_heading(doc, "18.2 2 分钟版本", 2)
    add_para(doc, "业务上，中台知识分散在产品文档、Git、Swagger、日志和指标平台，其他部门经常依赖人工对接。我把这些来源统一成一个只读问答入口。入口先用规则和轻量模型识别审批流、工作流、指标平台或 Bug；明确单领域直接进入专家，跨领域再由 Manager 汇总。RAG 链路用 Query Rewrite、字段化 BM25 和 pgvector 并行召回，RRF 融合后由 Qwen Rerank 精排，并用证据门禁区分代码存在、接口契约和部署状态。Bug 场景使用 LangGraph，要求环境和 trace ID，先查 Grafana 日志，再映射 develop/master 代码，支持 24 小时暂停恢复。工程上已经迁移到 PostgreSQL + pgvector，并建设了会话、长期记忆、SSE、引用详情和质量评测。当前能够用于 dev 小范围试用，但 Critical 还没有全部通过，下一步重点是领域证据补齐、P90 和多副本架构。")
    add_heading(doc, "18.3 5 分钟讲解顺序", 2)
    add_numbers(doc, [
        "用一个真实问题开场，例如“管理员转办接口如何对接”或“带 trace 的工作流异常”。",
        "说明为什么单纯大模型或纯向量不够：内部事实、精确标识符、环境分支和权限边界。",
        "画出入口—路由—专家—证据工具—混合检索—证据门禁—回答—评测的主链路。",
        "重点展开一个设计亮点：RAG 融合、LangGraph Bug、MCP 状态机或记忆隔离。",
        "说明 PostgreSQL + pgvector 迁移和性能取舍。",
        "用 Critical、引用覆盖、工具预算和 P50/P90 说明如何验证，而不是只展示 Demo。",
        "主动说明当前限制和下一步，体现工程判断。",
    ])
    add_heading(doc, "18.4 可量化表达", 2)
    add_bullets(doc, [
        "知识向量规模约 4 万条，而不是笼统说“很多数据”。",
        "单领域底层检索预算目标不超过 4 次，公开引用默认保留 3 至 5 条强证据。",
        "Critical 集为 30 条高风险用例，当前应如实说明已知 5 条基线仍为 3/5。",
        "目标体验可表述为网页单领域 P50 ≤ 15 秒、P90 ≤ 30 秒，但只有实测达标后才能说已经实现。",
    ])

    add_heading(doc, "19. 术语表", 1)
    add_table(doc, ["术语", "简明解释"], [
        ["RAG", "先检索外部知识，再让模型基于证据生成回答"],
        ["Embedding", "把文本映射为向量，用距离表示语义相似度"],
        ["BM25", "基于词频、逆文档频率和长度归一化的关键词排序算法"],
        ["pgvector", "PostgreSQL 的向量类型、距离算子与 ANN 索引扩展"],
        ["HNSW", "基于分层可导航小世界图的近似最近邻索引"],
        ["RRF", "用排名而非原始分数融合多个召回列表"],
        ["Rerank", "对少量候选做 query-document 联合相关性判断"],
        ["Agent", "能根据目标调用受控工具并组织结果的模型运行单元"],
        ["ReAct", "交替进行推理、行动和观察的 Agent 模式"],
        ["LangGraph", "用状态图构建可暂停、恢复和持久化的 Agent 工作流"],
        ["MCP", "让模型通过标准协议发现和调用外部工具/资源"],
        ["Checkpoint", "保存 Graph 结构化状态，以便跨轮恢复"],
        ["Citation", "支撑回答事实的文档、代码、接口或工具来源"],
        ["Grounding", "让模型结论受外部证据约束，而非只靠参数记忆"],
        ["LLM-as-a-Judge", "使用另一模型按 rubric 评价回答语义质量"],
        ["Critical", "带事实、工具、引用、安全和性能约束的高风险回归集"],
        ["Readiness", "服务是否具备接收流量所需关键依赖的状态检查"],
        ["SKIP LOCKED", "PostgreSQL 并发任务领取时跳过已被其他事务锁定的行"],
    ], [2300, 7060])

    add_heading(doc, "结语：如何真正掌握这个项目", 1)
    add_para(doc, "掌握项目的标准不是能背出技术栈，而是能回答四个问题：它解决了什么业务问题；一次请求经过哪些受控步骤；每个关键设计为什么这样取舍；如何用数据证明回答正确且足够快。建议以一条审批流接口问题和一条带 trace 的 Bug 为主线，亲自跟踪路由、检索、证据、回答和评测记录。能够解释失败案例并提出可验证改进，比只展示成功 Demo 更有说服力。")


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_toc(doc)
    build_content(doc)
    doc.core_properties.title = "中台知识问答 Agent 系统学习与面试指南"
    doc.core_properties.subject = "业务、架构、RAG、多智能体、LangGraph、记忆、评测与面试"
    doc.core_properties.author = "中台 Agent 项目组"
    doc.core_properties.keywords = "RAG, Agent, LangGraph, pgvector, 企业中台, 面试"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
