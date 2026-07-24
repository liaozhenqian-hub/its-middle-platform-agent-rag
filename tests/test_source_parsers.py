import json
from pathlib import Path
import stat
import subprocess
import warnings
import zipfile

from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject
import pytest

from knowledge.parsers.code import CodeParser, DirectoryDomainClassifier
from knowledge.parsers.documents import DocumentParser
from knowledge.parsers.policy import SourceFilePolicy
from knowledge.parsers.uploads import UnsafeArchiveError, extract_upload_archive
from knowledge.parsers.vue import VueSfcBatchParser, VueSfcSource


def _write_text_pdf(path: Path, text: str) -> None:
    _write_pdf_pages(path, [text])


def _write_pdf_pages(path: Path, pages: list[str | None]) -> None:
    writer = PdfWriter()
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_reference = writer._add_object(font)
    for text in pages:
        page = writer.add_blank_page(width=612, height=792)
        if text is None:
            continue
        page[NameObject("/Resources")] = DictionaryObject(
            {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_reference})}
        )
        commands = ["BT /F1 12 Tf 72 720 Td"]
        for index, line in enumerate(text.splitlines()):
            escaped = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
            if index:
                commands.append("0 -16 Td")
            commands.append(f"({escaped}) Tj")
        commands.append("ET")
        stream = DecodedStreamObject()
        stream.set_data(" ".join(commands).encode("ascii"))
        page[NameObject("/Contents")] = writer._add_object(stream)
    with path.open("wb") as output:
        writer.write(output)


def test_directory_classifier_uses_priority_rules_and_shared_fallback():
    classifier = DirectoryDomainClassifier(
        [
            ("**", "catch-all", 99),
            ("**/metric/**", "metric-platform", 10),
            ("**/approval/**", "approval-flow", 20),
            ("**/workflow/**", "workflow", 30),
        ]
    )

    assert classifier.classify("server/metric/service/A.java") == "metric-platform"
    assert classifier.classify("server/approval/A.java") == "approval-flow"
    assert classifier.classify("common/util/A.java") == "catch-all"


def test_source_file_policy_rejects_ignored_binary_unsupported_and_oversized_files():
    policy = SourceFilePolicy(max_file_bytes=16)

    assert policy.evaluate("src/Api.ts", b"export const x=1").accepted
    assert policy.evaluate(".git/config.java", b"class X {}").reason == "ignored_directory"
    assert policy.evaluate("node_modules/pkg/a.ts", b"const x=1").reason == "ignored_directory"
    assert policy.evaluate("target/generated/A.java", b"class A {}").reason == "ignored_directory"
    assert policy.evaluate("dist/a.js", b"const x=1").reason == "ignored_directory"
    assert policy.evaluate("build/a.js", b"const x=1").reason == "ignored_directory"
    assert policy.evaluate("src/logo.png", b"not-an-image").reason == "unsupported_extension"
    assert policy.evaluate("src/Binary.java", b"class X {}\x00").reason == "binary"
    assert policy.evaluate("src/Large.java", b"x" * 17).reason == "file_too_large"


def test_java_parser_emits_class_and_method_chunks_with_stable_identity():
    parser = CodeParser()
    source = """
package com.example.metric;
import java.util.List;

public class MetricService implements MetricQuery {
    public List<String> queryMetric(int appId) {
        return List.of("ok");
    }
}
"""

    first = parser.parse(
        relative_path="server/metric/MetricService.java",
        text=source,
        source_id="source-1",
        branch="main",
        commit_sha="abc123",
        domain_id="metric-platform",
    )
    second = parser.parse(
        relative_path="server/metric/MetricService.java",
        text=source.replace('"ok"', '"changed"'),
        source_id="source-1",
        branch="main",
        commit_sha="def456",
        domain_id="metric-platform",
    )

    symbols = {chunk.metadata["symbol_name"]: chunk for chunk in first}
    assert {"MetricService", "MetricService.queryMetric"} <= set(symbols)
    method = symbols["MetricService.queryMetric"]
    assert method.metadata["symbol_type"] == "method"
    assert method.metadata["start_line"] < method.metadata["end_line"]
    assert method.metadata["commit_sha"] == "abc123"
    assert "MetricService" in method.metadata["bm25_keywords"]
    second_method = next(
        chunk for chunk in second if chunk.metadata["symbol_name"] == "MetricService.queryMetric"
    )
    assert second_method.chunk_id == method.chunk_id


def test_java_parser_extracts_structure_calls_and_unique_overload_ids():
    parser = CodeParser()
    source = """
package com.example.metric;
import java.util.List;

@Deprecated
public class MetricService extends BaseService implements MetricQuery, Audited {
    @Inject private final MetricRepository repository;
    public MetricService(MetricRepository repository) { this.repository = repository; }
    public MetricService() { this(null); }
    @Override public String queryMetric(int id) { return repository.find(id); }
    public String queryMetric(String id) { return normalize(id); }
    class ResultMapper extends MapperBase implements Mapper {
        String map() { return queryMetric(1); }
    }
}
"""

    chunks = parser.parse(
        relative_path="server/metric/MetricService.java",
        text=source,
        source_id="source-1",
        branch="main",
        commit_sha="abc123",
        domain_id="metric-platform",
    )
    symbols = {}
    for chunk in chunks:
        symbols.setdefault(chunk.metadata["symbol_name"], []).append(chunk)

    assert len({chunk.chunk_id for chunk in chunks}) == len(chunks)
    assert symbols["MetricService.repository"][0].metadata["symbol_type"] == "field"
    service = symbols["MetricService"][0]
    assert json.loads(service.metadata["annotations"]) == ["Deprecated"]
    assert json.loads(service.metadata["extends"]) == ["BaseService"]
    assert json.loads(service.metadata["implements"]) == ["MetricQuery", "Audited"]
    int_overload, string_overload = symbols["MetricService.queryMetric"]
    assert int_overload.chunk_id != string_overload.chunk_id
    assert int_overload.metadata["signature"] != string_overload.metadata["signature"]
    assert json.loads(int_overload.metadata["calls"]) == ["repository.find"]
    assert json.loads(string_overload.metadata["calls"]) == ["normalize"]
    assert "MetricService.ResultMapper" in symbols
    assert "MetricService.ResultMapper.map" in symbols
    mapper = symbols["MetricService.ResultMapper"][0]
    assert json.loads(mapper.metadata["extends"]) == ["MapperBase"]
    assert json.loads(mapper.metadata["implements"]) == ["Mapper"]


def test_code_parser_scopes_nested_local_symbols_and_uses_block_identity():
    parser = CodeParser()
    java_chunks = parser.parse(
        relative_path="server/Scoped.java",
        text="""
class Scoped {
  void left() { class Local { void run() {} } }
  void right() { class Local { void run() {} } }
  void branch(boolean enabled) {
    if (enabled) { class BranchLocal {} }
    else { class BranchLocal {} }
  }
  void siblings(boolean first, boolean second) {
    if (first) { class SiblingLocal {} }
    if (second) { class SiblingLocal {} }
  }
}
""",
        source_id="source-1",
        branch="main",
        commit_sha="abc123",
        domain_id="shared",
    )
    java_names = [chunk.metadata["symbol_name"] for chunk in java_chunks]

    assert "Scoped.left.Local" in java_names
    assert "Scoped.right.Local" in java_names
    branch_locals = [
        chunk for chunk in java_chunks if chunk.metadata["symbol_name"] == "Scoped.branch.BranchLocal"
    ]
    assert len(branch_locals) == 2
    assert len({chunk.chunk_id for chunk in branch_locals}) == 2
    assert len({chunk.metadata["scope_identity"] for chunk in branch_locals}) == 2
    sibling_locals = [
        chunk
        for chunk in java_chunks
        if chunk.metadata["symbol_name"] == "Scoped.siblings.SiblingLocal"
    ]
    assert len(sibling_locals) == 2
    assert len({chunk.chunk_id for chunk in sibling_locals}) == 2
    assert len({chunk.metadata["scope_identity"] for chunk in sibling_locals}) == 2

    javascript_chunks = parser.parse(
        relative_path="web/scoped.js",
        text="""
export function left() { function local() { return load() } return local() }
export function right() { function local() { return load() } return local() }
""",
        source_id="source-1",
        branch="main",
        commit_sha="abc123",
        domain_id="shared",
    )
    javascript_names = {chunk.metadata["symbol_name"] for chunk in javascript_chunks}
    assert {"left.local", "right.local"} <= javascript_names


def test_code_parser_disambiguates_repeated_object_methods_with_same_signature():
    parser = CodeParser()
    chunks = parser.parse(
        relative_path="src/utils/directives.js",
        text="""
export default {
  install(app) {
    app.directive('first', { mounted(el) { el.focus() } })
    app.directive('second', { mounted(el) { el.select() } })
  }
}
""",
        source_id="source-1",
        branch="main",
        commit_sha="abc123",
        domain_id="shared",
    )

    mounted = [chunk for chunk in chunks if chunk.metadata["symbol_name"] == "install.mounted"]
    assert len(mounted) == 2
    assert len({chunk.chunk_id for chunk in mounted}) == 2


def test_code_parser_splits_oversized_symbols_into_stable_embedding_chunks():
    parser = CodeParser(max_chunk_chars=120)
    source = "class Large {\n  void execute() {\n" + "    callService();\n" * 30 + "  }\n}\n"

    first = parser.parse(
        relative_path="server/Large.java",
        text=source,
        source_id="source-1",
        branch="main",
        commit_sha="abc123",
        domain_id="shared",
    )
    second = parser.parse(
        relative_path="server/Large.java",
        text=source,
        source_id="source-1",
        branch="main",
        commit_sha="def456",
        domain_id="shared",
    )

    method_parts = [
        chunk for chunk in first if chunk.metadata["symbol_name"] == "Large.execute"
    ]
    second_parts = [
        chunk for chunk in second if chunk.metadata["symbol_name"] == "Large.execute"
    ]
    assert len(method_parts) > 1
    assert all(len(chunk.content) <= 120 for chunk in method_parts)
    assert len({chunk.chunk_id for chunk in method_parts}) == len(method_parts)
    assert [chunk.chunk_id for chunk in method_parts] == [
        chunk.chunk_id for chunk in second_parts
    ]
    assert [chunk.metadata["segment_index"] for chunk in method_parts] == list(
        range(len(method_parts))
    )
    assert all(
        chunk.metadata["segment_count"] == len(method_parts) for chunk in method_parts
    )


def test_typescript_parser_emits_interface_class_and_method():
    parser = CodeParser()
    chunks = parser.parse(
        relative_path="web/src/metric/api.ts",
        text="""
export interface Metric { id: number }
export class MetricApi {
  async queryMetric(id: number) { return fetch(`/metric/${id}`) }
}
export function formatMetric(value: number) { return `${value}` }
""",
        source_id="source-1",
        branch="main",
        commit_sha="abc123",
        domain_id="metric-platform",
    )

    names = {chunk.metadata["symbol_name"] for chunk in chunks}
    assert {"Metric", "MetricApi", "MetricApi.queryMetric", "formatMetric"} <= names


def test_typescript_parser_extracts_exports_arrow_functions_heritage_and_calls():
    parser = CodeParser()
    chunks = parser.parse(
        relative_path="web/src/metric/api.ts",
        text="""
import { request } from './request'
export type MetricId = string
export interface MetricClient { query(id: MetricId): Promise<void> }
export const loadMetric = (id: MetricId) => request(`/metrics/${id}`)
export default class MetricApi extends BaseApi implements MetricClient {
  query(id: MetricId) { return loadMetric(id) }
}
export { request }
""",
        source_id="source-1",
        branch="main",
        commit_sha="abc123",
        domain_id="metric-platform",
    )

    symbols = {chunk.metadata["symbol_name"]: chunk for chunk in chunks}
    assert "loadMetric" in symbols
    assert symbols["loadMetric"].metadata["symbol_type"] == "function"
    assert json.loads(symbols["loadMetric"].metadata["calls"]) == ["request"]
    assert json.loads(symbols["MetricApi"].metadata["extends"]) == ["BaseApi"]
    assert json.loads(symbols["MetricApi"].metadata["implements"]) == ["MetricClient"]
    assert {"MetricId", "MetricClient", "loadMetric", "MetricApi", "request"} <= set(
        json.loads(symbols["MetricApi"].metadata["exports"])
    )
    assert "MetricClient.query" in symbols


def test_typescript_parser_emits_interface_members_and_overload_signatures():
    chunks = CodeParser().parse(
        relative_path="web/src/client.ts",
        text="""
export interface Client {
  query(id: string): string
}
export class ClientImpl implements Client {
  query(id: string): string
  query(id: number): string
  query(id: string | number) { return String(id) }
}
""",
        source_id="source-1",
        branch="main",
        commit_sha="abc123",
        domain_id="shared",
    )

    names = [chunk.metadata["symbol_name"] for chunk in chunks]
    assert "Client.query" in names
    overloads = [chunk for chunk in chunks if chunk.metadata["symbol_name"] == "ClientImpl.query"]
    assert len(overloads) == 3
    assert len({chunk.metadata["signature"] for chunk in overloads}) == 3
    assert len({chunk.chunk_id for chunk in overloads}) == 3


def test_typescript_callable_signature_includes_return_type_generics_and_modifiers():
    chunks = CodeParser().parse(
        relative_path="web/src/signatures.ts",
        text="""
export interface Mapper {
  map<T>(value: string): Promise<T>
  map<T>(value: string): T
}
export abstract class BaseMapper {
  public abstract load<T>(id: string): Promise<T>
  protected static parse<U>(id: string): U
}
""",
        source_id="source-1",
        branch="main",
        commit_sha="abc123",
        domain_id="shared",
    )
    by_name = {}
    for chunk in chunks:
        by_name.setdefault(chunk.metadata["symbol_name"], []).append(chunk)

    map_overloads = by_name["Mapper.map"]
    assert len(map_overloads) == 2
    assert len({chunk.chunk_id for chunk in map_overloads}) == 2
    assert {chunk.metadata["signature"] for chunk in map_overloads} == {
        "<T> (value: string) : Promise<T>",
        "<T> (value: string) : T",
    }
    assert by_name["BaseMapper.load"][0].metadata["signature"] == (
        "public abstract <T> (id: string) : Promise<T>"
    )
    assert by_name["BaseMapper.parse"][0].metadata["signature"] == (
        "protected static <U> (id: string) : U"
    )


def test_vue_sfc_batch_parser_uses_one_process_and_preserves_paths_and_line_offsets(tmp_path: Path):
    calls = []

    def runner(command, **kwargs):
        calls.append((command, kwargs))
        request = json.loads(kwargs["input"])
        assert [item["relative_path"] for item in request["files"]] == [
            "web/src/metric/MetricPanel.vue",
            "web/src/workflow/WorkflowPanel.vue",
        ]
        return subprocess.CompletedProcess(
            command,
            0,
            stdout=json.dumps(
                {
                    "files": [
                        {
                            "relative_path": "web/src/metric/MetricPanel.vue",
                            "blocks": [
                                {
                                    "kind": "script",
                                    "language": "ts",
                                    "content": "export class MetricPanel { load() { return fetchMetric() } }",
                                    "start_line": 4,
                                }
                            ],
                        },
                        {
                            "relative_path": "web/src/workflow/WorkflowPanel.vue",
                            "blocks": [
                                {
                                    "kind": "script_setup",
                                    "language": "ts",
                                    "content": "export const loadWorkflow = () => requestWorkflow()",
                                    "start_line": 8,
                                }
                            ],
                        },
                    ]
                }
            ),
            stderr="",
        )

    parser = VueSfcBatchParser(helper_path=tmp_path / "parse-vue-sfc.mjs", runner=runner)
    chunks = parser.parse_many(
        [
            VueSfcSource(
                relative_path="web/src/metric/MetricPanel.vue",
                text="<script lang='ts'>...</script>",
                source_id="source-1",
                branch="main",
                commit_sha="abc123",
                domain_id="metric-platform",
            ),
            VueSfcSource(
                relative_path="web/src/workflow/WorkflowPanel.vue",
                text="<script setup lang='ts'>...</script>",
                source_id="source-1",
                branch="main",
                commit_sha="abc123",
                domain_id="workflow",
            ),
        ]
    )

    assert len(calls) == 1
    assert calls[0][0][0] == "node"
    symbols = {chunk.metadata["symbol_name"]: chunk for chunk in chunks}
    assert symbols["MetricPanel"].metadata["relative_path"] == "web/src/metric/MetricPanel.vue"
    assert symbols["MetricPanel"].metadata["start_line"] == 4
    assert symbols["MetricPanel.load"].metadata["start_line"] == 4
    assert symbols["loadWorkflow"].metadata["relative_path"] == (
        "web/src/workflow/WorkflowPanel.vue"
    )
    assert symbols["loadWorkflow"].metadata["start_line"] == 8


def test_vue_sfc_batch_parser_uses_block_kind_in_chunk_identity(tmp_path: Path):
    def runner(command, **kwargs):
        return subprocess.CompletedProcess(
            command,
            0,
            stdout=json.dumps(
                {
                    "files": [
                        {
                            "relative_path": "web/src/Duplicate.vue",
                            "blocks": [
                                {
                                    "kind": "script",
                                    "language": "ts",
                                    "content": "export const load = () => request()",
                                    "start_line": 2,
                                },
                                {
                                    "kind": "script_setup",
                                    "language": "ts",
                                    "content": "export const load = () => request()",
                                    "start_line": 6,
                                },
                            ],
                        }
                    ]
                }
            ),
            stderr="",
        )

    chunks = VueSfcBatchParser(
        helper_path=tmp_path / "parse-vue-sfc.mjs",
        runner=runner,
    ).parse_many(
        [
            VueSfcSource(
                relative_path="web/src/Duplicate.vue",
                text="<script>...</script><script setup>...</script>",
                source_id="source-1",
                branch="main",
                commit_sha="abc123",
                domain_id="shared",
            )
        ]
    )

    assert [chunk.metadata["symbol_name"] for chunk in chunks] == ["load", "load"]
    assert len({chunk.chunk_id for chunk in chunks}) == 2
    assert {chunk.metadata["scope_identity"] for chunk in chunks} == {
        "vue:script",
        "vue:script_setup",
    }


def test_document_parser_supports_markdown_txt_docx_and_blank_pdf(tmp_path: Path):
    markdown = tmp_path / "guide.md"
    markdown.write_text("# 指标平台\n概览\n## 查询指标\n使用应用 ID 查询。", encoding="utf-8")
    text = tmp_path / "notes.txt"
    text.write_text("审批流说明", encoding="utf-8")
    docx_path = tmp_path / "manual.docx"
    document = Document()
    document.add_heading("工作流", level=1)
    document.add_paragraph("创建流程实例。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "含义"
    document.save(docx_path)
    pdf_path = tmp_path / "scan.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with pdf_path.open("wb") as stream:
        writer.write(stream)

    parser = DocumentParser(max_chunk_chars=500)
    md_chunks = parser.parse(markdown, "source-md", "v1", "metric-platform")
    txt_chunks = parser.parse(text, "source-txt", "v1", "approval-flow")
    docx_chunks = parser.parse(docx_path, "source-docx", "v1", "workflow")
    pdf_chunks = parser.parse(pdf_path, "source-pdf", "v1", "workflow")

    assert [chunk.heading for chunk in md_chunks] == ["指标平台", "查询指标"]
    assert txt_chunks[0].content == "审批流说明"
    assert "创建流程实例" in docx_chunks[0].content
    assert "字段 | 含义" in docx_chunks[0].content
    assert pdf_chunks == []


def test_document_parser_preserves_supplied_relative_path_in_metadata_and_identity(tmp_path: Path):
    first = tmp_path / "frontend" / "readme.md"
    second = tmp_path / "backend" / "readme.md"
    first.parent.mkdir()
    second.parent.mkdir()
    first.write_text("# API\n接口说明", encoding="utf-8")
    second.write_text("# API\n接口说明", encoding="utf-8")
    parser = DocumentParser(max_chunk_chars=500)

    first_chunk = parser.parse(
        first,
        "source-docs",
        "v1",
        "metric-platform",
        relative_path="frontend/readme.md",
    )[0]
    second_chunk = parser.parse(
        second,
        "source-docs",
        "v1",
        "metric-platform",
        relative_path="backend/readme.md",
    )[0]

    assert first_chunk.metadata["relative_path"] == "frontend/readme.md"
    assert second_chunk.metadata["relative_path"] == "backend/readme.md"
    assert first_chunk.chunk_id != second_chunk.chunk_id


def test_document_parser_reports_ocr_required_and_sets_pdf_page_number(tmp_path: Path):
    text_pdf = tmp_path / "guide.pdf"
    _write_text_pdf(text_pdf, "Metric API Guide")
    blank_pdf = tmp_path / "scan.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with blank_pdf.open("wb") as output:
        writer.write(output)
    parser = DocumentParser(max_chunk_chars=500)

    parsed = parser.parse_with_diagnostics(
        text_pdf,
        "source-pdf",
        "v1",
        "metric-platform",
        relative_path="manuals/guide.pdf",
    )
    scanned = parser.parse_with_diagnostics(
        blank_pdf,
        "source-scan",
        "v1",
        "metric-platform",
        relative_path="manuals/scan.pdf",
    )

    assert parsed.diagnostics == []
    assert parsed.chunks[0].heading == "Metric API Guide"
    assert parsed.chunks[0].metadata["page_number"] == 1
    assert parsed.chunks[0].metadata["relative_path"] == "manuals/guide.pdf"
    assert scanned.chunks == []
    assert [diagnostic.code for diagnostic in scanned.diagnostics] == ["ocr_required"]


def test_document_parser_reports_blank_pages_in_mixed_pdf(tmp_path: Path):
    mixed_pdf = tmp_path / "mixed.pdf"
    _write_pdf_pages(mixed_pdf, ["Metric API Guide\nEndpoint reference", None])

    parsed = DocumentParser(max_chunk_chars=500).parse_with_diagnostics(
        mixed_pdf,
        "source-pdf",
        "v1",
        "metric-platform",
        relative_path="manuals/mixed.pdf",
    )

    assert [(chunk.heading, chunk.metadata["page_number"]) for chunk in parsed.chunks] == [
        ("Metric API Guide", 1)
    ]
    assert [(item.code, item.page_number) for item in parsed.diagnostics] == [
        ("ocr_required", 2)
    ]


def test_document_parser_preserves_docx_paragraph_table_order_and_heading_ownership(
    tmp_path: Path,
):
    path = tmp_path / "ordered.docx"
    document = Document()
    document.add_heading("First section", level=1)
    document.add_paragraph("Before first table")
    first_table = document.add_table(rows=1, cols=2)
    first_table.cell(0, 0).text = "First field"
    first_table.cell(0, 1).text = "First value"
    document.add_paragraph("After first table")
    document.add_heading("Second section", level=1)
    second_table = document.add_table(rows=1, cols=2)
    second_table.cell(0, 0).text = "Second field"
    second_table.cell(0, 1).text = "Second value"
    document.save(path)

    chunks = DocumentParser(max_chunk_chars=1000).parse(
        path,
        "source-docx",
        "v1",
        "workflow",
        relative_path="manuals/ordered.docx",
    )
    by_heading = {chunk.heading: chunk.content for chunk in chunks}

    first = by_heading["First section"]
    assert first.index("Before first table") < first.index("First field | First value")
    assert first.index("First field | First value") < first.index("After first table")
    assert "Second field" not in first
    assert "Second field | Second value" in by_heading["Second section"]


def test_upload_archive_rejects_path_traversal_and_extracts_relative_files(tmp_path: Path):
    safe_zip = tmp_path / "safe.zip"
    with zipfile.ZipFile(safe_zip, "w") as archive:
        archive.writestr("docs/readme.md", "ok")
    destination = tmp_path / "safe"

    extracted = extract_upload_archive(safe_zip, destination, max_files=10, max_bytes=100)

    assert extracted == [destination / "docs" / "readme.md"]

    unsafe_zip = tmp_path / "unsafe.zip"
    with zipfile.ZipFile(unsafe_zip, "w") as archive:
        archive.writestr("../secret.txt", "bad")
    try:
        extract_upload_archive(unsafe_zip, tmp_path / "unsafe", 10, 100)
    except UnsafeArchiveError:
        pass
    else:
        raise AssertionError("path traversal must be rejected")


def test_upload_archive_preflights_every_member_before_writing(tmp_path: Path):
    archive_path = tmp_path / "mixed.zip"
    with zipfile.ZipFile(archive_path, "w") as archive:
        archive.writestr("docs/readme.md", "safe")
        archive.writestr("../outside.md", "unsafe")
    destination = tmp_path / "extracted"

    with pytest.raises(UnsafeArchiveError, match="path traversal"):
        extract_upload_archive(archive_path, destination, max_files=10, max_bytes=100)

    assert not (destination / "docs" / "readme.md").exists()


@pytest.mark.parametrize(
    "member_name",
    [
        "/absolute/readme.md",
        "C:/absolute/readme.md",
        "C:\\absolute\\readme.md",
        "C:relative/readme.md",
    ],
)
def test_upload_archive_rejects_posix_and_windows_absolute_paths(
    tmp_path: Path,
    member_name: str,
):
    archive_path = tmp_path / "absolute.zip"
    with zipfile.ZipFile(archive_path, "w") as archive:
        archive.writestr(member_name, "unsafe")

    with pytest.raises(UnsafeArchiveError, match="absolute|traversal"):
        extract_upload_archive(archive_path, tmp_path / "destination", 10, 100)


def test_upload_archive_rejects_symbolic_links(tmp_path: Path):
    archive_path = tmp_path / "symlink.zip"
    link = zipfile.ZipInfo("docs/link.md")
    link.create_system = 3
    link.external_attr = (stat.S_IFLNK | 0o777) << 16
    with zipfile.ZipFile(archive_path, "w") as archive:
        archive.writestr(link, "../secret.md")

    with pytest.raises(UnsafeArchiveError, match="symbolic links"):
        extract_upload_archive(archive_path, tmp_path / "destination", 10, 100)


def test_upload_archive_enforces_per_file_limit_and_default_document_allowlist(tmp_path: Path):
    oversized = tmp_path / "oversized.zip"
    with zipfile.ZipFile(oversized, "w") as archive:
        archive.writestr("docs/large.txt", "x" * 11)
    with pytest.raises(UnsafeArchiveError, match="single file"):
        extract_upload_archive(
            oversized,
            tmp_path / "oversized",
            max_files=10,
            max_bytes=100,
            max_file_bytes=10,
        )

    mixed = tmp_path / "allowlist.zip"
    with zipfile.ZipFile(mixed, "w") as archive:
        archive.writestr("docs/readme.md", "knowledge")
        archive.writestr("docs/tool.exe", "binary")
    destination = tmp_path / "allowlist"

    extracted = extract_upload_archive(mixed, destination, max_files=10, max_bytes=100)

    assert extracted == [destination / "docs" / "readme.md"]
    assert not (destination / "docs" / "tool.exe").exists()


@pytest.mark.parametrize(
    "entries",
    [
        [("docs/readme.md", "first"), ("docs/readme.md", "second")],
        [("Docs/Readme.md", "first"), ("docs/readme.md", "second")],
        [("docs.md", "parent file"), ("docs.md/readme.md", "child")],
        [("docs.md/readme.md", "child"), ("docs.md", "parent file")],
    ],
)
def test_upload_archive_rejects_duplicate_case_and_file_parent_conflicts_before_writes(
    tmp_path: Path,
    entries: list[tuple[str, str]],
):
    archive_path = tmp_path / "conflict.zip"
    with warnings.catch_warnings():
        warnings.simplefilter("ignore", UserWarning)
        with zipfile.ZipFile(archive_path, "w") as archive:
            for name, content in entries:
                archive.writestr(name, content)
    destination = tmp_path / "destination"

    with pytest.raises(UnsafeArchiveError, match="duplicate|case-insensitive|parent"):
        extract_upload_archive(archive_path, destination, max_files=10, max_bytes=1000)

    assert not destination.exists()
